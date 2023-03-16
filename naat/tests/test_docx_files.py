from docx import Document
from naat_data import DOCX_DEJUSTICIA_PATH
import html

# https://python-docx.readthedocs.io/en/latest/user/quickstart.html
# docx_file.core_properties
# len(docx_file.sections)
# len(docx_file.tables)
# len(docx_file.paragraphs)
# add test with various docx


def test_docx_preprocessing():
    document = Document(DOCX_DEJUSTICIA_PATH)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text = run.text.replace("\xa0", " ")
            print(text)
